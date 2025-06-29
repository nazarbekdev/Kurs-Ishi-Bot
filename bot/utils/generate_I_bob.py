from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml.etree import fromstring
from dotenv import load_dotenv
import os
import re

# Atrof-muhit o‘zgaruvchilarini yuklash
load_dotenv()

# OpenAI mijozini boshlash
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# API ishlayotganligini tekshirish (test)
try:
    test_response = client.chat.completions.create(
        model="gpt-4.1-nano",
        messages=[{"role": "user", "content": "Salom, bu test xabar."}],
        temperature=0.7,
        max_tokens=50,
    )
    print(f"API test muvaffaqiyatli: {test_response.choices[0].message.content}")
except Exception as e:
    print(f"API xatosi: {str(e)}")
    exit(1)

# LaTeX formulalarini aniqlash uchun regex
latex_pattern = r'(\$\$[\s\S]*?\$\$|\\\[[\s\S]*?\\\]|\\\(.*?\\\)|\\begin\{cases\}[\s\S]*?\\end\{cases\})'


def clean_generated_text(text: str, til: str) -> str:
    unwanted_sections = {
        "o'zbek tili": [
            r"(?i)^#+?\s*Kirish\b.*?(?=(?:^#+?\s*[A-Za-z0-9])|\Z)",
            r"(?i)^#+?\s*Xulosa\b.*?(?=(?:^#+?\s*[A-Za-z0-9])|\Z)",
            r"(?i)^#+?\s*Foydalanilgan adabiyotlar\b.*?(?=(?:^#+?\s*[A-Za-z0-9])|\Z)",
            r"(?i)^Kirish\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^Xulosa\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^Foydalanilgan adabiyotlar\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^\*\*Xulosa\*\*\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^__Xulosa__\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^\*Xulosa\*\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^_Xulosa_\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^Xulosa\s*[:–-].*?(?=\n[A-Za-z0-9]|$)"
        ],
        "rus tili": [
            r"(?i)^#+?\s*Введение\b.*?(?=(?:^#+?\s*[A-Za-z0-9])|\Z)",
            r"(?i)^#+?\s*Заключение\b.*?(?=(?:^#+?\s*[A-Za-z0-9])|\Z)",
            r"(?i)^#+?\s*Список литературы\b.*?(?=(?:^#+?\s*[A-Za-z0-9])|\Z)",
            r"(?i)^Введение\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^Заключение\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^Список литературы\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^\*\*Заключение\*\*\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^__Заключение__\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^\*Заключение\*\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^_Заключение_\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^Заключение\s*[:–-].*?(?=\n[A-Za-z0-9]|$)"
        ],
        "ingliz tili": [
            r"(?i)^#+?\s*Introduction\b.*?(?=(?:^#+?\s*[A-Za-z0-9])|\Z)",
            r"(?i)^#+?\s*Conclusion\b.*?(?=(?:^#+?\s*[A-Za-z0-9])|\Z)",
            r"(?i)^#+?\s*References\b.*?(?=(?:^#+?\s*[A-Za-z0-9])|\Z)",
            r"(?i)^Introduction\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^Conclusion\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^References\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^\*\*Conclusion\*\*\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^__Conclusion__\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^\*Conclusion\*\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^_Conclusion_\b.*?(?=\n[A-Za-z0-9]|$)",
            r"(?i)^Conclusion\s*[:–-].*?(?=\n[A-Za-z0-9]|$)"
        ]
    }

    lines = text.split('\n')
    cleaned_lines = []
    skip_section = False

    for line in lines:
        if not line.strip() or line.strip().startswith(("```", "|", "- ", "* ", "> ", "#")):
            if not skip_section:
                cleaned_lines.append(line)
            continue

        section_found = False
        for pattern in unwanted_sections.get(til.lower(), []):
            if re.match(pattern, line, re.MULTILINE | re.DOTALL):
                skip_section = True
                section_found = True
                break

        if not section_found and line.strip().startswith("#"):
            skip_section = False

        if not skip_section:
            cleaned_lines.append(line)

    final_lines = []
    skip_rest = False
    for line in cleaned_lines:
        if skip_rest:
            continue
        if re.search(r"(?i)\bXulosa\b", line):
            skip_rest = True
            continue
        final_lines.append(line)

    return '\n'.join(final_lines).strip()


def calculate_word_and_token_count(total_pages: int, num_sections: int = 3) -> tuple[int, int]:
    if total_pages < 11:
        words_per_page = 500
        tokens_per_section = 1500
    elif 10 < total_pages < 21:
        words_per_page = 1000
        tokens_per_section = 3000
    elif 20 < total_pages < 31:
        words_per_page = 2000
        tokens_per_section = 6000
    elif 30 < total_pages < 41:
        words_per_page = 2500
        tokens_per_section = 8000
    elif 40 < total_pages < 70:
        words_per_page = 3000
        tokens_per_section = 9000

    return words_per_page, tokens_per_section


PROMPT_TEMPLATES = {
    "o'zbek tili": {
        "section_prompt": lambda fan_nomi, mavzu, title, num: (
            f"'{fan_nomi}' fanidan '{mavzu}' mavzusida faqat '{title}' bo‘limi haqida ilmiy ma'lumot ber, "
            f"{num} so‘zdan iborat bo‘lsin, O‘zbek tilida, ilmiy uslubda, plagiatdan xoli bo‘lsin. "
            "Faqat bir bo‘limning mazmunini bering, boshqa hech qanday qo‘shimcha bo‘lim yoki umumiy tuzilma qo‘shmang! "
            "Hech qanday Kirish, Xulosa yoki Foydalanilgan adabiyotlar qismlarini qo‘shmang! "
            "Agar xulosa yoki boshqa bo‘lim qo‘shmoqchi bo‘lsang, o‘rniga '{title}' bo‘limi mazmunini davom ettir! "
            "Matnni to‘g‘ridan-to‘g‘ri bo‘lim mazmuni sifatida boshlang, sarlavhadan keyin darhol kontentni bering. "
            "Bo‘limni yozishda '{title}' bo‘limining asosiy nuqtalariga e’tibor qaratib, ular bo‘yicha chuqur ma’lumot ber, "
            "ortiqcha subheadings (###) qo‘shishdan saqlan, faqat zarur bo‘lganda minimal subheadings ishlat. "
            "Ma’lumotni iloji boricha chuqur yorit, umumiy ma’lumotlardan ko‘ra aniq faktlar, tahlillar va misollar keltir. Times New Roman, 14 pt."
            "Matematik formulalar bo‘lsa, ularni LaTeX formatida (\$$ ...\ $$, \$$ ...\ $$, \\begin{{cases}}...\\end{{cases}}) aniq va to‘g‘ri yoz. "
            "Iloji bo'lsa Markdown formatida ro‘yxatlar (- yoki 1.), jadvallar (| bilan) va kod misollari (``` bilan) qo‘shing."
        )
    },
    "rus tili": {
        "section_prompt": lambda fan_nomi, mavzu, title, num: (
            f"'{fan_nomi}' дисциплины по теме '{mavzu}' подробно опиши только раздел '{title}' на русском языке, "
            f"текст должен состоять из {num} слов, в научном стиле, без плагиата. "
            "Дай только содержимое одного раздела, не добавляй другие разделы или общую структуру! "
            "Не включай Введение, Заключение или Список литературы! "
            "Если хочешь добавить заключение или другой раздел, вместо этого продолжи содержание раздела '{title}'! "
            "Начни текст сразу с содержания раздела, после заголовка сразу давай контент. "
            "При написании раздела сосредоточься на ключевых аспектах раздела '{title}', предоставь глубокую информацию по этим аспектам, "
            "избегай добавления лишних подзаголовков (###), используй минимальное количество подзаголовков только при необходимости. "
            "Старайся дать глубокое освещение темы, приводи конкретные факты, анализы и примеры, а не общую информацию. Times New Roman, 14 pt."
            "Если есть математические формулы, пиши их в формате LaTeX (\\[...\\], \\(...\\), \\begin{{cases}}...\\end{{cases}}) четко и правильно. "
            "По возможности используй Markdown-форматирование: списки (- или 1.), таблицы (| с разделителями) и примеры кода (```)."
        )
    },
    "ingliz tili": {
        "section_prompt": lambda fan_nomi, mavzu, title, num: (
            f"In the field of '{fan_nomi}' under the topic '{mavzu}', provide a detailed scientific explanation "
            f"of only the section '{title}' in English, consisting of {num} words, in a scientific style, free of plagiarism. "
            "Provide only the content of one section, do not add any other sections or overall structure! "
            "Do not include Introduction, Conclusion, or References sections! "
            "If you are tempted to add a conclusion or another section, instead continue the content of the current section '{title}'! "
            "Start the text directly with the section content, immediately after the heading, provide the content. "
            "While writing the section, focus on the key aspects of the section '{title}', provide in-depth information on these aspects, "
            "avoid adding unnecessary subheadings (###), and use minimal subheadings only when necessary. "
            "Aim to provide deep coverage of the topic, including specific facts, analyses, and examples, rather than general information. Times New Roman, 14 pt."
            "If there are mathematical equations, write them clearly and correctly in LaTeX format (\$$ ...\ $$, \$$ ...\ $$, \\begin{{cases}}...\\end{{cases}}). "
            "If possible, use Markdown formatting: lists (- or 1.), tables (with | separators), and code examples (```)."
        )
    }
}


def sanitize_filename(filename: str) -> str:
    return re.sub(r'[^\w\s-]', '', filename).strip().replace(' ', '_')


def latex_to_omml(latex: str) -> str:
    latex = latex.strip()
    return '<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:oMath><m:r><m:t>' + latex + '</m:t></m:r></m:oMath></m:oMathPara>'


def add_equation(document: Document, latex: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    omml_str = latex_to_omml(latex)
    omml_element = fromstring(omml_str)
    paragraph._element.append(omml_element)


def add_formatted_paragraph(document: Document, text: str, is_heading: bool = False, is_center: bool = False,
                            is_bold: bool = False, is_italic: bool = False) -> None:
    parts = re.split(latex_pattern, text)
    para = document.add_paragraph()

    for part in parts:
        if not part:
            continue

        if re.match(latex_pattern, part):
            add_equation(document, part)
            para = document.add_paragraph()
            continue

        if is_heading:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if is_center else WD_PARAGRAPH_ALIGNMENT.LEFT
            run = para.add_run(part.strip())
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
        else:
            if part.startswith("#"):
                level = 0
                for char in part:
                    if char == "#":
                        level += 1
                    else:
                        break
                cleaned_text = part[level:].strip()
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run = para.add_run(cleaned_text)
                run.bold = True
                run.font.name = 'Times New Roman'
                if level == 1:
                    run.font.size = Pt(18)
                    run.font.name = 'Times New Roman'
                elif level == 2:
                    run.font.size = Pt(16)
                    run.font.name = 'Times New Roman'
                else:
                    run.font.size = Pt(14)
                    run.font.name = 'Times New Roman'
                if level >= 4:
                    para.paragraph_format.left_indent = Pt(10)
            elif part.startswith("- ") or part.startswith("* ") or bool(re.match(r'^\d+\.\s', part)):
                if part.startswith("- ") or part.startswith("* "):
                    para.style = 'List Bullet'
                    cleaned_text = part[2:].strip()
                else:
                    para.style = 'List Number'
                    cleaned_text = re.sub(r'^\d+\.\s', '', part).strip()
                sub_parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_|~~.*?~~)', cleaned_text)
                for sub_part in sub_parts:
                    if sub_part.startswith("**") and sub_part.endswith("**") or sub_part.startswith(
                            "__") and sub_part.endswith("__"):
                        run = para.add_run(sub_part[2:-2].strip())
                        run.bold = True
                        run.font.size = Pt(14)
                        run.font.name = 'Times New Roman'
                    elif sub_part.startswith("*") and sub_part.endswith("*") or sub_part.startswith(
                            "_") and sub_part.endswith("_"):
                        run = para.add_run(sub_part[1:-1].strip())
                        run.italic = True
                        run.font.size = Pt(14)
                        run.font.name = 'Times New Roman'
                    elif sub_part.startswith("~~") and sub_part.endswith("~~"):
                        run = para.add_run(sub_part[2:-2].strip())
                        run.font.strike = True
                        run.font.size = Pt(14)
                        run.font.name = 'Times New Roman'
                    else:
                        run = para.add_run(sub_part.strip())
                        run.font.size = Pt(14)
                        run.font.name = 'Times New Roman'
            elif part.startswith("> "):
                para.paragraph_format.left_indent = Pt(20)
                cleaned_text = part[2:].strip()
                run = para.add_run(cleaned_text)
                run.italic = True
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'
            elif part.startswith("`") and part.endswith("`"):
                cleaned_text = part[1:-1].strip()
                run = para.add_run(cleaned_text)
                run.font.name = "Courier New"
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.font.size = Pt(14)
            else:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                sub_parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_|~~.*?~~|`.*?`)', part)
                for sub_part in sub_parts:
                    if sub_part.startswith("**") and sub_part.endswith("**") or sub_part.startswith(
                            "__") and sub_part.endswith("__"):
                        run = para.add_run(sub_part[2:-2].strip())
                        run.bold = True
                        run.font.size = Pt(14)
                        run.font.name = "Times New Roman"
                    elif sub_part.startswith("*") and sub_part.endswith("*") or sub_part.startswith(
                            "_") and sub_part.endswith("_"):
                        run = para.add_run(sub_part[1:-1].strip())
                        run.italic = True
                        run.font.size = Pt(14)
                        run.font.name = "Times New Roman"
                    elif sub_part.startswith("~~") and sub_part.endswith("~~"):
                        run = para.add_run(sub_part[2:-2].strip())
                        run.font.strike = True
                        run.font.size = Pt(14)
                        run.font.name = "Times New Roman"
                    elif sub_part.startswith("`") and sub_part.endswith("`"):
                        run = para.add_run(sub_part[1:-1].strip())
                        run.font.name = "Courier New"
                        run.font.color.rgb = RGBColor(128, 128, 128)
                        run.font.size = Pt(14)
                    else:
                        run = para.add_run(sub_part.strip())
                        run.bold = is_bold
                        run.italic = is_italic
                        run.font.size = Pt(14)
                        run.font.name = "Times New Roman"

    para.paragraph_format.space_after = Pt(8)


def add_table(document: Document, lines: list) -> None:
    table_rows = [line.strip() for line in lines if line.strip().startswith("|") and line.strip().endswith("|")]
    if len(table_rows) < 2:
        return

    headers = [header.strip() for header in table_rows[0].strip("|").split("|") if header.strip()]
    data_rows = [[cell.strip() for cell in row.strip("|").split("|") if cell.strip()] for row in table_rows[2:]]

    table = document.add_table(rows=len(data_rows) + 1, cols=len(headers))
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(14)

    for row_idx, row in enumerate(data_rows):
        for col_idx, cell_text in enumerate(row):
            cell = table.cell(row_idx + 1, col_idx)
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.size = Pt(14)


def add_code_block(document: Document, lines: list) -> None:
    for line in lines:
        para = document.add_paragraph()
        para.paragraph_format.left_indent = Pt(20)
        run = para.add_run(line.rstrip())
        run.font.name = "Courier New"
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.size = Pt(14)
        para.paragraph_format.space_after = Pt(8)


def generate_bob_1(fan_nomi: str, mavzu: str, til: str, chapter_1_sections: dict, total_pages: int) -> str:
    if not chapter_1_sections:
        chapter_1_sections = {
            "chapter_title": "I bob. Noma'lum",
            "1.1.": "Bo‘lim 1.1",
            "1.2.": "Bo‘lim 1.2",
            "1.3.": "Bo‘lim 1.3"
        }

    words_per_section, tokens_per_section = calculate_word_and_token_count(total_pages)

    template = PROMPT_TEMPLATES.get(til.lower(), PROMPT_TEMPLATES["o'zbek tili"])

    document = Document()

    chapter_title = chapter_1_sections.get("chapter_title", "I. BOB. NOMA'LUM")
    add_formatted_paragraph(document, chapter_title, is_heading=True, is_center=True)

    for section_key in ["1.1.", "1.2.", "1.3."]:
        section_title = f"{section_key} {chapter_1_sections.get(section_key, 'Nomalum bolim')}"
        add_formatted_paragraph(document, section_title, is_heading=True, is_center=False)

        try:
            prompt = template["section_prompt"](fan_nomi, mavzu, chapter_1_sections.get(section_key, section_key),
                                                words_per_section)
            print(f"Prompt (I bob, {section_key}): {prompt}")
            response = client.chat.completions.create(
                model="gpt-4.1-nano",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
            )
            section_text = response.choices[0].message.content.strip()
            print(f"Generatsiya qilingan matn (I bob, {section_key}): {section_text[:200]}...")

            section_text = clean_generated_text(section_text, til)
            print(f"Tozalangan matn (I bob, {section_key}): {section_text[:200]}...")
        except Exception as e:
            print(f"OpenAI xatosi (I bob, {section_key}): {str(e)}")
            section_text = f"{section_title} bo‘yicha ma'lumotlar hozircha mavjud emas."

        lines = section_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            if line.startswith("```"):
                code_block_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_block_lines.append(lines[i])
                    i += 1
                add_code_block(document, code_block_lines)
                i += 1 if i < len(lines) else 0
                continue

            if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                add_table(document, table_lines)
                continue

            add_formatted_paragraph(document, line, is_heading=False)
            i += 1

    try:
        sanitized_mavzu = sanitize_filename(mavzu)
        save_path = f"generated_docs/bob_1_{sanitized_mavzu}_{til.lower().replace(' ', '_')}.docx"
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        document.save(save_path)
        print(f"Fayl saqlangan: {save_path}")
    except Exception as e:
        raise Exception(f"Fayl saqlashda xato: {str(e)}")

    return save_path


def generate_bob_2(fan_nomi: str, mavzu: str, til: str, chapter_2_sections: dict, total_pages: int) -> str:
    if not chapter_2_sections:
        chapter_2_sections = {
            "chapter_title": "II bob. Noma'lum",
            "2.1.": "Bo‘lim 2.1",
            "2.2.": "Bo‘lim 2.2",
            "2.3.": "Bo‘lim 2.3"
        }

    words_per_section, tokens_per_section = calculate_word_and_token_count(total_pages)

    template = PROMPT_TEMPLATES.get(til.lower(), PROMPT_TEMPLATES["o'zbek tili"])

    document = Document()

    chapter_title = chapter_2_sections.get("chapter_title", "II. BOB. NOMA'LUM")
    add_formatted_paragraph(document, chapter_title, is_heading=True, is_center=True)

    for section_key in ["2.1.", "2.2.", "2.3."]:
        section_title = f"{section_key} {chapter_2_sections.get(section_key, 'Nomalum bolim')}"
        add_formatted_paragraph(document, section_title, is_heading=True, is_center=False)

        try:
            prompt = template["section_prompt"](fan_nomi, mavzu, chapter_2_sections.get(section_key, section_key),
                                                words_per_section)
            print(f"Prompt (II bob, {section_key}): {prompt}")
            response = client.chat.completions.create(
                model="gpt-4.1-nano",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
            )
            section_text = response.choices[0].message.content.strip()
            print(f"Generatsiya qilingan matn (II bob, {section_key}): {section_text[:200]}...")

            section_text = clean_generated_text(section_text, til)
            print(f"Tozalangan matn (II bob, {section_key}): {section_text[:200]}...")
        except Exception as e:
            print(f"OpenAI xatosi (II bob, {section_key}): {str(e)}")
            section_text = f"{section_title} bo‘yicha ma'lumotlar hozircha mavjud emas."

        lines = section_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            if line.startswith("```"):
                code_block_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_block_lines.append(lines[i])
                    i += 1
                add_code_block(document, code_block_lines)
                i += 1 if i < len(lines) else 0
                continue

            if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                add_table(document, table_lines)
                continue

            add_formatted_paragraph(document, line, is_heading=False)
            i += 1

    try:
        sanitized_mavzu = sanitize_filename(mavzu)
        save_path = f"generated_docs/bob_2_{sanitized_mavzu}_{til.lower().replace(' ', '_')}.docx"
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        document.save(save_path)
        print(f"Fayl saqlangan: {save_path}")
    except Exception as e:
        raise Exception(f"Fayl saqlashda xato: {str(e)}")

    return save_path


if __name__ == "__main__":
    fan_nomi = "Fizika"
    mavzu = "Elektromagnit Maydonlar"
    til = "o'zbek tili"
    chapter_1_sections = {
        'chapter_title': 'I bob. Elektromagnit Maydonlarning Nazariy Asoslari',
        '1.1.': 'Maxvell Tenglamalari va Ularning Ahamiyati',
        '1.2.': 'Elektromagnit Maydonlarni Modellash Usullari',
        '1.3.': 'Elektromagnit To‘lqinlarning Tarqalishi'
    }
    chapter_2_sections = {
        'chapter_title': 'II bob. Elektromagnit Maydonlarning Amaliy Tadbiqlari',
        '2.1.': 'Elektromagnit Maydonlarning Texnologiyalarda Qo‘llanilishi',
        '2.2.': 'Elektromagnit To‘lqinlarning Tibbiyotdagi Roli',
        '2.3.': 'Elektromagnit Maydonlarning Kelajakdagi Istiqbollari'
    }
    total_pages = 20

    try:
        bob_1_path = generate_bob_1(fan_nomi, mavzu, til, chapter_1_sections, total_pages)
        print(f"I bob fayli muvaffaqiyatli yaratildi: {bob_1_path}")

        bob_2_path = generate_bob_2(fan_nomi, mavzu, til, chapter_2_sections, total_pages)
        print(f"II bob fayli muvaffaqiyatli yaratildi: {bob_2_path}")
    except Exception as e:
        print(f"Xatolik yuz berdi: {str(e)}")
