from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import logging

# Logger sozlash
logger = logging.getLogger(__name__)


def create_coursework_docx(content, topic):
    doc = Document()

    # Sahifa chegaralarini sozlash
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Sarlavha
    heading = doc.add_heading(f"KURS ISHI\nMavzu: {topic}", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    # Mundarija
    doc.add_heading("MUNDARIJA", level=1)

    # Matn bo‘sh bo‘lsa
    if not content or not content.strip():
        logger.error("DOCX generatsiyasi: Matn bo‘sh yoki mavjud emas")
        p = doc.add_paragraph("Xatolik: Kurs ishi matni generatsiya qilinmadi.")
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
        file_path = f"media/courseworks/{topic}.docx"
        doc.save(file_path)
        return file_path

    # Matnni bo‘limlarga ajratish
    sections = {
        "KIRISH": "",
        "1-BO‘LIM": "",
        "2-BO‘LIM": "",
        "3-BO‘LIM": "",
        "XULOSA": "",
        "FOYDALANILGAN ADABIYOTLAR RO‘YXATI": ""
    }

    current_section = None
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue

        clean_line = line.replace('#', '').strip()
        line_upper = clean_line.upper()

        if "KIRISH" in line_upper:
            current_section = "KIRISH"
        elif "I BO‘LIM" in line_upper or "NAZARIY TAHLIL" in line_upper:
            current_section = "1-BO‘LIM"
        elif "II BO‘LIM" in line_upper or "AMALIY QISM" in line_upper:
            current_section = "2-BO‘LIM"
        elif "III BO‘LIM" in line_upper or "TAKLIFLAR" in line_upper:
            current_section = "3-BO‘LIM"
        elif "XULOSA" in line_upper:
            current_section = "XULOSA"
        elif "FOYDALANILGAN ADABIYOTLAR" in line_upper:
            current_section = "FOYDALANILGAN ADABIYOTLAR RO‘YXATI"
        else:
            if current_section:
                sections[current_section] += line + "\n"

    # Bo‘limlarni hujjatga yozish
    for section_title, section_content in sections.items():
        if section_content.strip():
            heading = doc.add_heading(section_title, level=1)
            for run in heading.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

            # Matn ichidagi kod, formula va jadval ajratish
            paragraphs = section_content.split('\n')
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue

                # Kod blok
                if para.startswith("```") and para.endswith("```"):
                    code_text = para.strip("```")
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(12)
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.line_spacing = 1.5
                # Formula
                elif para.startswith('$$') and para.endswith('$$'):
                    formula = para.strip('$$')
                    p = doc.add_paragraph()
                    run = p.add_run(formula)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.bold = True
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.line_spacing = 1.5
                # Jadval
                elif '|' in para and '-' in para:
                    rows = [r.split('|')[1:-1] for r in para.split('\n') if '|' in r]
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        for i, row in enumerate(rows):
                            for j, cell_text in enumerate(row):
                                table.cell(i, j).text = cell_text.strip()
                else:
                    p = doc.add_paragraph(para)
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                    p.paragraph_format.line_spacing = 1.5

    # Faylni saqlash
    file_path = f"media/courseworks/{topic}.docx"
    doc.save(file_path)
    logger.info(f"DOCX fayl saqlandi: {file_path}")
    return file_path
