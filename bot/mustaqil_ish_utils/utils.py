from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import os


def add_page_numbers(document: Document) -> None:
    """
    Hujjatning har bir sahifasiga pastki qismda sahifa raqamlarini qo'shadi.

    Args:
        document (Document): Sahifa raqamlari qo'shiladigan docx hujjat.
    """
    # Hujjatning birinchi bo'limini olish (agar bo'limlar bo'lmasa, yangi bo'lim avtomatik yaratiladi)
    section = document.sections[0]

    # Footer (pastki qism) ni olish
    footer = section.footer

    # Yangi paragraf qo'shish (footer ichida)
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Raqamni o'rtada joylashtirish

    # Sahifa raqamini qo'shish uchun Open XML elementlaridan foydalanamiz
    run = para.add_run()

    # Sahifa raqami uchun field qo'shish (PAGE)
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char1)

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    run._r.append(instr_text)

    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char2)

    # Formatlash: shrift o'lchami va qalinlik
    run.font.size = Pt(10)  # Kichikroq shrift o'lchami
    run.bold = False


def copy_paragraph_formatting(from_para, to_para):
    """
    Paragrafning formatlashini (masalan, tekislanish, bo'shliqlar) ko'chirish.
    """
    to_para.alignment = from_para.alignment
    to_para.paragraph_format.space_before = from_para.paragraph_format.space_before
    to_para.paragraph_format.space_after = from_para.paragraph_format.space_after
    to_para.paragraph_format.line_spacing = from_para.paragraph_format.line_spacing
    to_para.paragraph_format.left_indent = from_para.paragraph_format.left_indent
    to_para.paragraph_format.right_indent = from_para.paragraph_format.right_indent
    to_para.paragraph_format.first_line_indent = from_para.paragraph_format.first_line_indent


def copy_run_formatting(from_run, to_run):
    """
    Run (matn qismi) formatlashini (masalan, qalin, shrift o'lchami) ko'chirish.
    """
    to_run.bold = from_run.bold
    to_run.italic = from_run.italic
    to_run.underline = from_run.underline
    to_run.font.size = from_run.font.size
    to_run.font.name = from_run.font.name
    to_run.font.color.rgb = from_run.font.color.rgb
    to_run.font.all_caps = from_run.font.all_caps
    to_run.font.strike = from_run.font.strike


def merge_docs(output_path: str, file_paths: list) -> str:
    """
    Bir nechta docx fayllarni birlashtirib, bitta docx fayl sifatida saqlaydi.
    Har bir docx faylning tarkibi yangi sahifadan boshlanadi va formatlash saqlanadi.
    Oxirida sahifa raqamlari qo'shiladi.

    Args:
        output_path (str): Natija faylning saqlanadigan yo'li.
        file_paths (list): Birlashtiriladigan docx fayllar yo'llari ro'yxati.

    Returns:
        str: Birlashtirilgan docx faylning yo'li.

    Raises:
        FileNotFoundError: Agar fayl topilmasa xato chiqaradi.
    """
    # Yangi docx hujjat yaratish
    merged_document = Document()

    for i, docx_path in enumerate(file_paths):
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Fayl topilmadi: {docx_path}")

        # Agar bu birinchi docx fayl bo'lmasa, yangi sahifa qo'shish
        if i > 0:
            merged_document.add_page_break()

        # Docx faylni o'qish
        doc = Document(docx_path)

        # Har bir paragrafni ko'chirish
        for para in doc.paragraphs:
            if not para.text.strip():  # Bo'sh paragraflarni o'tkazib yuborish
                continue

            # Yangi paragraf yaratish
            new_para = merged_document.add_paragraph()

            # Paragraf formatlashini ko'chirish
            copy_paragraph_formatting(para, new_para)

            # Har bir run (matn qismi) ni ko'chirish
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                copy_run_formatting(run, new_run)

    # Sahifa raqamlarini qo'shish
    add_page_numbers(merged_document)

    # Birlashtirilgan faylni saqlash
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    merged_document.save(output_path)
    print(f"Birlashtirilgan fayl saqlandi: {output_path}")
    return output_path


def add_formatted_paragraph(document: Document, text: str, is_heading: bool = False, is_center: bool = False,
                            is_bold: bool = False) -> None:
    para = document.add_paragraph()
    if is_center:
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run = para.add_run(text.strip())
    if is_bold or is_heading:
        run.bold = True
    run.font.size = Pt(14)
    para.paragraph_format.space_after = Pt(8)
