from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
import os
import re

# Atrof-muhit o‘zgaruvchilarini yuklash
load_dotenv()

# OpenAI mijozini boshlash
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# "Foydalanilgan adabiyotlar" sarlavhasini tilga moslashtirish uchun lug‘at
REFERENCES_TITLES = {
    "o'zbek tili": "Foydalanilgan adabiyotlar",
    "rus tili": "Список литературы",
    "ingliz tili": "References"
}

# Prompt uchun tilga mos shablonlar
PROMPT_TEMPLATES = {
    "o'zbek tili": {
        "references": "Foydalanilgan adabiyotlar",
        "format": "GOST standarti",
        "example": """
        Ahmedov A.A. Yurak kasalliklari va ularning oldini olish. Toshkent: Fan, 2020. 150 bet.
        Saidova M.B. Kardiologiyada zamonaviy yondashuvlar // Tibbiyot jurnali. 2021. №3. 25-30-betlar.
        World Health Organization. Cardiovascular Diseases [Internet]. 2023. URL: https://www.who.int/health-topics/cardiovascular-diseases (Murojaat sanasi: 2025-yil 1-may).
        """
    },
    "rus tili": {
        "references": "Список литературы",
        "format": "GOST standarti",
        "example": """
        Ахмедов А.А. Сердечные заболевания и их профилактика. Ташкент: Фан, 2020. 150 с.
        Саидова М.Б. Современные подходы в кардиологии // Журнал медицины. 2021. №3. С. 25-30.
        World Health Organization. Cardiovascular Diseases [Интернет]. 2023. URL: https://www.who.int/health-topics/cardiovascular-diseases (Дата обращения: 1 мая 2025 г.).
        """
    },
    "ingliz tili": {
        "references": "References",
        "format": "APA standarti",
        "example": """
        Ahmedov, A. A. (2020). Heart diseases and their prevention. Fan.
        Saidova, M. B. (2021). Modern approaches in cardiology. Journal of Medicine, (3), 25-30.
        World Health Organization. (2023). Cardiovascular diseases. https://www.who.int/health-topics/cardiovascular-diseases
        """
    }
}


def sanitize_filename(filename: str) -> str:
    """Fayl nomini maxsus belgilardan tozalash."""
    return re.sub(r'[^\w\s-]', '', filename).strip().replace(' ', '_')


def add_formatted_paragraph(document: Document, text: str, is_heading: bool = False) -> None:
    """Matnni formatlangan holda DOCX fayliga qo‘shish."""
    para = document.add_paragraph()

    # Matnni formatlash
    if is_heading:
        # Sarlavha uchun markazda va qalin
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(text.strip())
        run.bold = True
        run.font.size = Pt(14)
    else:
        # Oddiy matn uchun chapdan va qalin emas
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(text.strip())
        run.bold = False
        run.font.size = Pt(14)

    para.paragraph_format.space_after = Pt(8)


def calculate_references_count(page_count: int) -> int:
    """Sahifalar soniga qarab manbalar sonini hisoblash."""
    # Har 5 sahifa uchun 1-2 manba deb hisoblaymiz
    base_references = page_count // 3
    total_references = max(5, base_references + 5)  # Kamida 5 manba bo‘ladi
    return min(total_references, 30)  # Maksimal 30 manba bilan cheklaymiz


def generate_foydalanilgan_adabiyotlar(fan_nomi: str, mavzu: str, til: str, page_count: int, reja_items: list) -> str:
    """
    Foydalanilgan adabiyotlar qismini generatsiya qilish va DOCX fayl sifatida saqlash.

    Args:
        fan_nomi (str): Fanning nomi.
        mavzu (str): Mustaqil ish mavzusi.
        til (str): Til ("o'zbek tili", "rus tili", "ingliz tili").
        page_count (int): Mustaqil ishning sahifalar soni.
        reja_items (list): Asosiy qism bo‘limlari ro‘yxati.
        client (OpenAI): OpenAI mijoz obyekti.

    Returns:
        str: Saqlangan DOCX fayl yo'li.
    """

    client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

    # 1. Manbalar sonini hisoblash
    references_count = calculate_references_count(page_count)

    # Tilga mos shablonni olish
    template = PROMPT_TEMPLATES.get(til.lower(), PROMPT_TEMPLATES["o'zbek tili"])  # Default: O‘zbek tili

    # Reja_items ni promptga qo‘shish uchun tayyorlash
    reja_info = "\n".join([f"{item}" for item in reja_items])

    # 2. OpenAI'dan manbalar ro‘yxatini olish
    if til.lower() == "ingliz tili":
        prompt = f"""
        '{mavzu}' mavzusida mustaqil ish uchun '{template['references']}' qismini {template['format']} bo‘yicha yozib ber. Quyidagi talablarga rioya qil:
        - Jami {references_count} ta manba bo‘lsin.
        - Manbalar turli xil bo‘lsin: kitoblar, ilmiy maqolalar, veb-saytlar (Internet manbalari).
        - Har bir manba {template['format']} bo‘yicha formatda bo‘lsin:
          - Kitoblar uchun: Author, A. A. (Year). Title of the book. Publisher.
          - Maqolalar uchun: Author, A. A. (Year). Title of the article. Journal Name, (Issue), pages.
          - Internet manbalari uchun: Author/Organization. (Year). Title of the page. URL
        - Manbalar ro‘yxati tartiblangan holda (alfavitik tartibda) bo‘lsin.
        - Har bir manba oldiga raqam qo‘shmang, faqat manba matnini yozing.
        - Quyidagi bo‘limlar asosida manbalar mos bo‘lsin:
          Asosiy qism bo‘limlari:
          {reja_info}
        - Matn ilmiy uslubda, {til} tilida bo‘lsin.
        Namuna sifatida quyidagi uslubdan foydalaning:
        {template['example']}
        """
    else:
        prompt = f"""
        '{mavzu}' mavzusida mustaqil ish uchun '{template['references']}' qismini {template['format']} asosida yozib ber. Quyidagi talablarga rioya qil:
        - Jami {references_count} ta manba bo‘lsin.
        - Manbalar turli xil bo‘lsin: kitoblar, ilmiy maqolalar, veb-saytlar (Internet manbalari).
        - Har bir manba {template['format']} bo‘yicha formatda bo‘lsin:
          - Kitoblar uchun: Muallif. Kitob nomi. Nashr joyi: Nashriyot, yil. Sahifalar soni.
          - Maqolalar uchun: Muallif. Maqola nomi // Jurnal nomi. Yil. №. Sahifalar.
          - Internet manbalari uchun: Tashkilot yoki muallif. Material nomi [Internet]. Yil. URL: link (Murojaat sanasi: 2025-yil 1-may).
        - Manbalar ro‘yxati tartiblangan holda (raqamli tartibda) bo‘lsin.
        - Har bir manba oldiga raqam qo‘shmang, faqat manba matnini yozing.
        - Quyidagi bo‘limlar asosida manbalar mos bo‘lsin:
          Asosiy qism bo‘limlari:
          {reja_info}
        - Matn ilmiy uslubda, {til} tilida bo‘lsin.
        Namuna sifatida quyidagi uslubdan foydalaning:
        {template['example']}
        """

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-nano",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=1500
        )
        adabiyotlar_text = response.choices[0].message.content.strip()
        print(f"Generatsiya qilingan adabiyotlar: {adabiyotlar_text[:200]}...")
    except Exception as e:
        print(f"OpenAI xatosi (Adabiyotlar): {str(e)}")
        adabiyotlar_text = "1. Adabiyotlar hozircha mavjud emas."

    # 3. DOCX fayl yaratish
    document = Document()

    # Foydalanilgan adabiyotlar sarlavhasi (markazda, qalin, tilga mos)
    references_title = REFERENCES_TITLES.get(til.lower(), "Foydalanilgan adabiyotlar")
    add_formatted_paragraph(document, references_title, is_heading=True)

    # Manbalar ro‘yxatini qo‘shish
    lines = adabiyotlar_text.split('\n')
    cleaned_lines = []
    for line in lines:
        if not line.strip():
            continue
        # Boshidagi raqamlarni olib tashlash (agar mavjud bo‘lsa)
        cleaned_line = re.sub(r'^\d+\.\s*', '', line.strip())
        if cleaned_line:
            cleaned_lines.append(cleaned_line)

    # Har bir manba oldiga raqam qo‘shish (1., 2., ...)
    for i, line in enumerate(cleaned_lines, 1):
        formatted_line = f"{i}. {line}"
        add_formatted_paragraph(document, formatted_line, is_heading=False)

    # 4. Saqlash
    try:
        sanitized_mavzu = sanitize_filename(mavzu)
        save_path = f"generated_docs/foydalanilgan_adabiyotlar_{sanitized_mavzu}_{til.lower().replace(' ', '_')}.docx"
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        document.save(save_path)
        print(f"Foydalanilgan adabiyotlar fayli saqlandi: {save_path}")
    except Exception as e:
        raise Exception(f"Fayl saqlashda xato: {str(e)}")

    return save_path


if __name__ == "__main__":
    from dotenv import load_dotenv

    load_dotenv()
    fan_nomi = "Pedagogika"
    mavzu = "MAKTABGACHA TA’LIM MUASSASASI, OILA VA MAKTAB HAMKORLIGI"
    til = "o'zbek tili"
    page_count = 20
    reja_items = [
        "1. Maktabgacha ta’lim tashkilotining oila bilan hamkorligi mazmuni.",
        "2. Ota-onalar burchagi va ko‘rgazma tashkil etish.",
        "3. Ota-onalar qo‘mitasi.",
        "4. Aholi o‘rtasida pedagogik bilimlarni targ‘ib qilish.",
        "5. Bolalarni maktabga tayyorlashda maktab bilan hamkorlikning ahamiyati."
    ]
    save_path = generate_foydalanilgan_adabiyotlar(fan_nomi, mavzu, til, page_count, reja_items)
