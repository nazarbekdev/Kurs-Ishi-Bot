from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
import os
import re

# Atrof-muhit o‘zgaruvchilarini yuklash
load_dotenv()

# OpenAI mijozini boshlash
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# "MUNDARIJA" so‘zini tilga moslashtirish uchun lug‘at
TOC_TITLES = {
    "o'zbek tili": "MUNDARIJA",
    "rus tili": "СОДЕРЖАНИЕ",
    "ingliz tili": "TABLE OF CONTENTS"
}

# Har bir til uchun kalit so‘zlarni aniqlash (qalin qilish uchun)
BOLD_KEYWORDS = {
    "o'zbek tili": ['kirish', 'bob', 'xulosa', 'foydalanilgan adabiyotlar'],
    "rus tili": ['введение', 'глава', 'заключение', 'список литературы'],
    "ingliz tili": ['introduction', 'chapter', 'conclusion', 'references']
}

# Prompt uchun tilga mos shablonlar
PROMPT_TEMPLATES = {
    "o'zbek tili": {
        "introduction": "Kirish",
        "chapter": "bob",
        "subchapter": "kichik mavzu",
        "conclusion": "Xulosa",
        "references": "Foydalanilgan adabiyotlar"
    },
    "rus tili": {
        "introduction": "Введение",
        "chapter": "глава",
        "subchapter": "подраздел",
        "conclusion": "Заключение",
        "references": "Список литературы"
    },
    "ingliz tili": {
        "introduction": "Introduction",
        "chapter": "chapter",
        "subchapter": "subsection",
        "conclusion": "Conclusion",
        "references": "References"
    }
}


def sanitize_filename(filename: str) -> str:
    """Fayl nomini maxsus belgilardan tozalash."""
    return re.sub(r'[^\w\s-]', '', filename).strip().replace(' ', '_')


def generate_mundarija(fan_nomi: str, mavzu: str, til: str) -> tuple[str, str, dict, dict]:
    """
    Mundarija generatsiya qilish, docx fayl yaratish va I va II bob bo'limlarini qaytarish.

    Args:
        fan_nomi (str): Fanning nomi.
        mavzu (str): Kurs ishi mavzusi.
        til (str): Til ("o'zbek tili", "rus tili", "ingliz tili").

    Returns:
        tuple: (mundarija matni, docx fayl yo'li, I bob bo'limlari lug'ati, II bob bo'limlari lug'ati)
    """
    # Tilga mos shablonni olish
    template = PROMPT_TEMPLATES.get(til.lower(), PROMPT_TEMPLATES["o'zbek tili"])  # Default: O‘zbek tili

    # 1. OpenAI'dan mundarija uchun matn olish
    prompt = f"""
    '{fan_nomi}' fanidan '{mavzu}' mavzusida kurs ishi uchun mundarija tuzib ber. Barcha tuzilishi {til} da va quyidagi ko'rinishda bo'lsin:
     {template['introduction']}
     I-{template['chapter']} (nomi)
      1.1. {template['subchapter']}
      1.2. {template['subchapter']}
      1.3. {template['subchapter']}
     II-{template['chapter']} (nomi)
      2.1. {template['subchapter']}
      2.2. {template['subchapter']}
      2.3. {template['subchapter']}
     {template['conclusion']}
     {template['references']}
    Har bir {template['chapter']} va {template['subchapter']}lar mavzuga mos holda, ilmiy ohangda bo'lsin. 
    Bo'limlar oldida "- " yoki boshqa belgilar bo'lmasin, faqat "1.1." formatida bo'lsin.
    Qo'yilgan talablardan tashqari ortiqcha yozuv va belgilar bo'lmasin!!!
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
        )
        mundarija_text = response.choices[0].message.content.strip()
    except Exception as e:
        raise Exception(f"OpenAI API xatosi: {str(e)}")

    # 2. I va II bob bo'limlarini ajratib olish
    chapter_1_sections = {}  # I bob bo'limlari
    chapter_2_sections = {}  # II bob bo'limlari
    current_chapter = None  # Hozirgi bobni kuzatish uchun (None, "chapter_1", "chapter_2")
    lines = mundarija_text.split('\n')
    formatted_mundarija_text = TOC_TITLES.get(til.lower(), "MUNDARIJA") + "\n\n"

    # Birinchi tekshirish: I va II bob bo'limlarini ajratish
    for line in lines:
        if not line.strip():
            continue
        formatted_line = line.strip()
        indent = 0

        # I-bob sarlavhasini aniqlash
        if (f"I-{template['chapter']}".lower() in formatted_line.lower() or
                formatted_line.lower().startswith("i ") or
                "chapter i" in formatted_line.lower()):
            current_chapter = "chapter_1"
            if "chapter_title" not in chapter_1_sections:  # Faqat birinchi marta saqlansin
                chapter_1_sections["chapter_title"] = formatted_line

        # II-bob sarlavhasini aniqlash
        elif (f"II-{template['chapter']}".lower() in formatted_line.lower() or
              formatted_line.lower().startswith("ii ") or
              "chapter ii" in formatted_line.lower()):
            current_chapter = "chapter_2"
            if "chapter_title" not in chapter_2_sections:  # Faqat birinchi marta saqlansin
                chapter_2_sections["chapter_title"] = formatted_line

        # I-bob bo'limlari (1.1, 1.2, 1.3)
        elif re.match(r'^1\.\d+\.', formatted_line) and current_chapter == "chapter_1":
            indent = 20
            section_key = formatted_line.split(" ")[0]  # 1.1, 1.2, 1.3
            section_title = " ".join(formatted_line.split(" ")[1:]).strip()
            chapter_1_sections[section_key] = section_title

        # II-bob bo'limlari (2.1, 2.2, 2.3)
        elif re.match(r'^2\.\d+\.', formatted_line) and current_chapter == "chapter_2":
            indent = 20
            section_key = formatted_line.split(" ")[0]  # 2.1, 2.2, 2.3
            section_title = " ".join(formatted_line.split(" ")[1:]).strip()
            chapter_2_sections[section_key] = section_title

        # Agar Xulosa yoki Foydalanilgan adabiyotlar bo'lsa, bob bo'limlarini tugatamiz
        elif formatted_line.lower() in [template['conclusion'].lower(), template['references'].lower()]:
            current_chapter = None

        # Mundarija matnini formatlash (foydalanuvchiga ko'rsatish uchun)
        indent_str = "  " * (indent // 10)
        formatted_mundarija_text += indent_str + formatted_line + "\n"

    # Agar II bob bo‘limlari bo‘sh bo‘lsa, OpenAI javobidan qayta tekshirish
    if len(chapter_2_sections) <= 1:  # Faqat chapter_title bo‘lsa, bo‘limlar yo‘q demakdir
        current_chapter = None
        chapter_2_sections = {}  # Qayta boshlash
        for line in lines:
            if not line.strip():
                continue
            formatted_line = line.strip()
            # II-bob sarlavhasini aniqlash
            if (f"II-{template['chapter']}".lower() in formatted_line.lower() or
                    formatted_line.lower().startswith("ii ") or
                    "chapter ii" in formatted_line.lower()):
                current_chapter = "chapter_2"
                if "chapter_title" not in chapter_2_sections:  # Faqat birinchi marta saqlansin
                    chapter_2_sections["chapter_title"] = formatted_line

            # II-bob bo'limlari (2.1, 2.2, 2.3)
            elif re.match(r'^2\.\d+\.', formatted_line) and current_chapter == "chapter_2":
                section_key = formatted_line.split(" ")[0]  # 2.1, 2.2, 2.3
                section_title = " ".join(formatted_line.split(" ")[1:]).strip()
                chapter_2_sections[section_key] = section_title

            # Agar Xulosa yoki Foydalanilgan adabiyotlar bo'lsa, bob bo'limlarini tugatamiz
            elif formatted_line.lower() in [template['conclusion'].lower(), template['references'].lower()]:
                current_chapter = None

        # Agar hali ham bo‘sh bo‘lsa, xatolik haqida xabar berish
        if len(chapter_2_sections) <= 1:
            raise ValueError("II bob bo‘limlari OpenAI javobida topilmadi. Iltimos, promptni qayta ko‘rib chiqing.")

    # 3. DOCX fayl yaratish
    document = Document()

    # Mavzu nomi (markazda, qalin)
    title = document.add_paragraph(mavzu.upper())
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(14)

    # MUNDARIJA sarlavhasi (tilga mos ravishda, markazda, qalin)
    toc_title = TOC_TITLES.get(til.lower(), "MUNDARIJA")
    mundarija_title = document.add_paragraph(toc_title)
    mundarija_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = mundarija_title.runs[0]
    run.bold = True
    run.font.size = Pt(14)

    # Mundarija matnini qo'shish (faqat OpenAI javobidan olingan qatorlarni ishlatamiz)
    bold_keywords = BOLD_KEYWORDS.get(til.lower(), BOLD_KEYWORDS["o'zbek tili"])
    for line in lines:
        if not line.strip():
            continue
        indent = 0
        formatted_line = line.strip()
        # "1." yoki "2." bilan boshlanadigan qatorlarni aniqlash va indent qo'shish
        if re.match(r'^(1|2)\.\d+\.', formatted_line):
            indent = 20
        para = document.add_paragraph(formatted_line)
        para.paragraph_format.left_indent = Pt(indent)
        run = para.runs[0]
        run.font.size = Pt(14)
        # Qalin qilish uchun kalit so‘zlarni tekshirish
        if any(keyword in formatted_line.lower() for keyword in bold_keywords):
            run.bold = True
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.paragraph_format.space_after = Pt(8)

    # 4. Saqlash
    sanitized_mavzu = sanitize_filename(mavzu)
    save_path = f"generated_docs/mundarija_{sanitized_mavzu}_{til.lower().replace(' ', '_')}.docx"
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    document.save(save_path)

    # Mundarija matni, fayl yo'li, I va II bob bo'limlarini qaytarish
    return formatted_mundarija_text, save_path, chapter_1_sections, chapter_2_sections


if __name__ == "__main__":
    # Test uchun har bir til uchun mavzuni moslashtirish
    mavzular = {
        # "o'zbek tili": "Yurak urishi va uni to'xtab qolish sabablari",
        # "rus tili": "Причины сердцебиения и его остановки",
        "ingliz tili": "Causes of Heartbeat and Its Cessation",
    }

    fan_nomi = "Tibbiyot"
    for til, mavzu in mavzular.items():
        print(f"{til} uchun mundarija yaratilmoqda...")
        try:
            mundarija_text, save_path, chapter_1_sections, chapter_2_sections = generate_mundarija(fan_nomi, mavzu, til)
            print(f"Mundarija matni:\n{mundarija_text}")
            print(f"I bob bo'limlari: {chapter_1_sections}")
            print(f"II bob bo'limlari: {chapter_2_sections}")
            print(f"Fayl muvaffaqiyatli yaratildi: {save_path}\n")
        except Exception as e:
            print(f"Xatolik yuz berdi ({til}): {str(e)}")
