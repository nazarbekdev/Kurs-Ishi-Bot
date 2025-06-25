from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
import os
import re

# Atrof-muhit o‘zgaruvchilarini yuklash
load_dotenv()

# OpenAI mijozini boshlash
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# "Kirish" sarlavhasini tilga moslashtirish uchun lug‘at
INTRODUCTION_TITLES = {
    "o'zbek tili": "Kirish",
    "rus tili": "Введение",
    "ingliz tili": "Introduction"
}

# Har bir til uchun kalit so‘zlarni aniqlash (qalin qilish uchun)
BOLD_KEYWORDS = {
    "o'zbek tili": ['Kurs ishining maqsadi va vazifalari', 'Kurs ishining predmeti', 'Kurs ishining obyekti'],
    "rus tili": ['Цели и задачи курсовой работы', 'Предмет курсовой работы', 'Объект курсовой работы'],
    "ingliz tili": ['Objectives and Tasks of the Course Work', 'Subject of the Course Work',
                    'Object of the Course Work']
}

# Prompt uchun tilga mos shablonlar
PROMPT_TEMPLATES = {
    "o'zbek tili": {
        "introduction": "Kirish",
        "objectives": "Kurs ishining maqsadi va vazifalari",
        "subject": "Kurs ishining predmeti",
        "object": "Kurs ishining obyekti",
        "objectives_start": "Ushbu kurs ishining asosiy maqsadi",
        "subject_start": "Kurs ishining predmeti",
        "object_start": "Kurs ishining obyekti",
        "example": """
        Bugungi kunda tibbiyot sohasida yurak urishining fiziologiyasi va uni to‘xtab qolish sabablarini o‘rganish dolzarb masalalardan biridir.
        Kurs ishining maqsadi va vazifalari: Ushbu kurs ishining asosiy maqsadi yurak urishining fiziologik mexanizmlarini va uni to‘xtab qolish sabablarini tahlil qilishdir.
        Kurs ishining predmeti: Kurs ishining predmeti yurak urishining fiziologik jarayonlari va ularning buzilishiga olib keluvchi omillardan iboratdir.
        Kurs ishining obyekti: Kurs ishining obyekti yurak-qon tomir tizimi va uning tashqi omillarga ta’siridir.
        """
    },
    "rus tili": {
        "introduction": "Введение",
        "objectives": "Цели и задачи курсовой работы",
        "subject": "Предмет курсовой работы",
        "object": "Объект курсовой работы",
        "objectives_start": "Основная цель данной курсовой работы заключается в",
        "subject_start": "Предметом курсовой работы является",
        "object_start": "Объектом курсовой работы является",
        "example": """
        В современном мире изучение физиологических процессов, таких как сердцебиение, играет ключевую роль в медицине и биологии.
        Цели и задачи курсовой работы: Основная цель данной курсовой работы заключается в исследовании физиологических механизмов сердцебиения и причин его остановки.
        Предмет курсовой работы: Предметом курсовой работы являются физиологические процессы, связанные с сердцебиением, и факторы, приводящие к его остановке.
        Объект курсовой работы: Объектом курсовой работы являются сердечно-сосудистая система и внешние факторы, влияющие на её функционирование.
        """
    },
    "ingliz tili": {
        "introduction": "Introduction",
        "objectives": "Objectives and Tasks of the Course Work",
        "subject": "Subject of the Course Work",
        "object": "Object of the Course Work",
        "objectives_start": "The main objective of this course work is to",
        "subject_start": "The subject of the course work is",
        "object_start": "The object of the course work is",
        "example": """
        In today's world, understanding physiological processes such as heartbeat is crucial in the fields of medicine and biology.
        Objectives and Tasks of the Course Work: The main objective of this course work is to investigate the physiological mechanisms of heartbeat and the reasons for its cessation.
        Subject of the Course Work: The subject of the course work is the physiological processes related to heartbeat and the factors leading to its cessation.
        Object of the Course Work: The object of the course work is the cardiovascular system and the external factors affecting its functioning.
        """
    }
}


def sanitize_filename(filename: str) -> str:
    """Fayl nomini maxsus belgilardan tozalash."""
    return re.sub(r'[^\w\s-]', '', filename).strip().replace(' ', '_')


def add_formatted_paragraph(document: Document, text: str, is_heading: bool = False, heading_text: str = None) -> None:
    """Matnni formatlangan holda DOCX fayliga qo‘shish, sarlavha qalin va xat boshidan, qolgan matn oddiy."""
    para = document.add_paragraph()

    if is_heading and heading_text:
        # Sarlavha bo‘lgan qatorni chap tomondan tekislaymiz
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Sarlavha va matnni ajratish
        if text.lower().startswith(heading_text.lower() + ":"):
            heading_part = heading_text + ":"
            content_part = text[len(heading_part):].strip()

            # Sarlavha qismini qalin qilish
            run = para.add_run(heading_part)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"

            # Qolgan matnni oddiy qilish
            run = para.add_run(" " + content_part)
            run.bold = False
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"
        else:
            # Agar sarlavha bo‘lmasa, qatorni ikki tomondan tekislaymiz
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            run = para.add_run(text.strip())
            run.bold = False
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"
    else:
        # Umumiy kirish matni yoki boshqa oddiy matnlar ikki tomondan tekislanadi
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        run = para.add_run(text.strip())
        run.bold = False
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"

    para.paragraph_format.space_after = Pt(8)


def generate_kirish(fan_nomi: str, mavzu: str, til: str) -> str:
    """
    Kirish qismini generatsiya qilib, docx fayl sifatida saqlaydi va yo'lini qaytaradi.

    Args:
        fan_nomi (str): Fan nomi.
        mavzu (str): Kurs ishi mavzusi.
        til (str): Til ("o'zbek tili", "rus tili", "ingliz tili").

    Returns:
        str: Saqlangan docx fayl yo'li.
    """
    # Tilga mos shablonni olish
    template = PROMPT_TEMPLATES.get(til.lower(), PROMPT_TEMPLATES["o'zbek tili"])  # Default: O‘zbek tili

    # 1. OpenAI'dan kirish qismi uchun matn olish
    prompt = f"""
    '{mavzu}' mavzusida kurs ishi uchun '{template['introduction']}' qismini yozib ber. Quyidagi tuzilma va uslubda bo‘lsin, {til} tilida:
    Sarlavhasiz, 500 ta so'z bilan mavzuning dolzarbligi va ahamiyatini ta'kidlang:
      Bunda mavzuning dolzarbligi, ahamiyati va o‘rganish zarurati haqida umumiy ma'lumot bering.
    - {template['objectives']} (birinchi jumla sarlavhadan keyin ikki nuqta bilan bir qatorda davom etsin, maqsad va vazifalarni aniq va ilmiy uslubda 500 ta so'z bilan ifodab yozing):
      Birinchi qator "{template['objectives']}: {template['objectives_start']} ..." shaklida bo‘lsin.
    - {template['subject']} (birinchi jumla sarlavhadan keyin ikki nuqta bilan bir qatorda davom etsin, predmetni aniq va ilmiy uslubda 500 ta so'z bilan ifodab yozing):
      Birinchi qator "{template['subject']}: {template['subject_start']} ..." shaklida bo‘lsin.
    - {template['object']} (birinchi jumla sarlavhadan keyin ikki nuqta bilan bir qatorda davom etsin, obyektni aniq va ilmiy uslubda 500 ta so'z bilan ifodab yozing):
      Birinchi qator "{template['object']}: {template['object_start']} ..." shaklida bo‘lsin.
    Har bir bo‘limni keng doirada fikrlab yoz. Matn ilmiy uslubda, aniq va ravon, plagiatdan xoli bo‘lsin. {til} tilida yozing. Times New Roman, 14 pt.
    Har bir bo‘limni yangi qatordan boshlang. Umumiy kirish matni uchun sarlavha ishlatmang, qolgan bo‘limlarning birinchi jumlasi sarlavhadan keyin ikki nuqta (:) bilan bir qatorda davom etsin.
    Namuna sifatida quyidagi uslubdan foydalaning:
    {template['example']}
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-nano",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
        )
        kirish_text = response.choices[0].message.content.strip()
    except Exception as e:
        raise Exception(f"OpenAI API xatosi: {str(e)}")

    # 2. DOCX fayl yaratish
    document = Document()

    # Kirish sarlavhasi (markazda, qalin, tilga mos)
    introduction_title = INTRODUCTION_TITLES.get(til.lower(), "Kirish")  # Agar til topilmasa, default O‘zbek tili
    kirish_title = document.add_paragraph(introduction_title)
    kirish_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = kirish_title.runs[0]
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    # Kirish matnini qo‘shish
    lines = kirish_text.split('\n')
    bold_keywords = BOLD_KEYWORDS.get(til.lower(),
                                      BOLD_KEYWORDS["o'zbek tili"])  # Agar til topilmasa, default O‘zbek tili

    current_section = None
    paragraph_buffer = []  # Paragrafni yig‘ish uchun bufer

    for line in lines:
        if not line.strip():
            continue
        # Sarlavha ekanligini aniqlash
        heading_match = None
        for keyword in bold_keywords:
            if line.strip().lower().startswith(keyword.lower() + ":"):
                heading_match = keyword
                break

        if heading_match:
            # Agar oldin yig‘ilgan paragraf bo‘lsa, uni qo‘shish
            if paragraph_buffer:
                add_formatted_paragraph(document, "\n".join(paragraph_buffer), is_heading=(current_section is not None),
                                        heading_text=current_section)
                paragraph_buffer = []
            current_section = heading_match
            paragraph_buffer.append(line.strip())
        else:
            # Matnni buferga qo‘shish
            paragraph_buffer.append(line.strip())

    # Oxirgi paragrafni qo‘shish
    if paragraph_buffer:
        add_formatted_paragraph(document, "\n".join(paragraph_buffer), is_heading=(current_section is not None),
                                heading_text=current_section)

    # 3. Saqlash
    try:
        sanitized_mavzu = sanitize_filename(mavzu)
        save_path = f"generated_docs/kirish_{sanitized_mavzu}_{til.lower().replace(' ', '_')}.docx"
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        document.save(save_path)
    except Exception as e:
        raise Exception(f"Fayl saqlashda xato: {str(e)}")

    return save_path
