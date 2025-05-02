from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from dotenv import load_dotenv
import os
import re

# Atrof-muhit o‘zgaruvchilarini yuklash
load_dotenv()

# OpenAI mijozini boshlash
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))


def sanitize_filename(filename: str) -> str:
    """Fayl nomini maxsus belgilardan tozalash."""
    return re.sub(r'[^\w\s-]', '', filename).strip().replace(' ', '_')


def add_formatted_paragraph(document: Document, text: str, is_heading: bool = False, is_center: bool = False,
                            is_bold: bool = False, is_italic: bool = False, is_list: str = None) -> None:
    """Matnni formatlangan holda DOCX fayliga qo‘shish, Markdown belgilarni hisobga olgan holda."""
    para = document.add_paragraph()

    # Matnni formatlash
    if is_heading:
        # Sarlavha uchun
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if is_center else WD_PARAGRAPH_ALIGNMENT.LEFT
        run = para.add_run(text.strip())
        run.bold = True
        run.font.size = Pt(14)
    else:
        # Markdown belgilarni tahlil qilish
        text = text.strip()

        # 1. Sarlavhalar (### yoki ####)
        if text.startswith("#"):
            level = 0
            for char in text:
                if char == "#":
                    level += 1
                else:
                    break
            cleaned_text = text[level:].strip()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = para.add_run(cleaned_text)
            run.bold = True
            # Darajaga qarab shrift hajmini o‘zgartirish (ixtiyoriy)
            if level == 1:
                run.font.size = Pt(18)
            elif level == 2:
                run.font.size = Pt(16)
            else:
                run.font.size = Pt(14)
            if level >= 4:
                para.paragraph_format.left_indent = Pt(10)  # Ichki sarlavhalar uchun indentatsiya
        # 2. Tartibli yoki tartibsiz ro‘yxat
        elif text.startswith("- ") or text.startswith("* ") or bool(re.match(r'^\d+\.\s', text)):
            if text.startswith("- ") or text.startswith("* "):
                para.style = 'List Bullet'
                cleaned_text = text[2:].strip()
            else:
                para.style = 'List Number'
                cleaned_text = re.sub(r'^\d+\.\s', '', text).strip()
            # Ichki formatlash (bold, italic)
            parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_|~~.*?~~)', cleaned_text)
            for part in parts:
                if part.startswith("**") and part.endswith("**") or part.startswith("__") and part.endswith("__"):
                    run = para.add_run(part[2:-2].strip())
                    run.bold = True
                    run.font.size = Pt(14)
                elif part.startswith("*") and part.endswith("*") or part.startswith("_") and part.endswith("_"):
                    run = para.add_run(part[1:-1].strip())
                    run.italic = True
                    run.font.size = Pt(14)
                elif part.startswith("~~") and part.endswith("~~"):
                    run = para.add_run(part[2:-2].strip())
                    run.font.strike = True
                    run.font.size = Pt(14)
                else:
                    run = para.add_run(part.strip())
                    run.font.size = Pt(14)
        # 3. Iqtiboslar
        elif text.startswith("> "):
            para.paragraph_format.left_indent = Pt(20)
            cleaned_text = text[2:].strip()
            run = para.add_run(cleaned_text)
            run.italic = True
            run.font.size = Pt(14)
        # 4. Inline kod
        elif text.startswith("`") and text.endswith("`"):
            cleaned_text = text[1:-1].strip()
            run = para.add_run(cleaned_text)
            run.font.name = "Courier New"
            run.font.color.rgb = RGBColor(128, 128, 128)  # Kulrang rang
            run.font.size = Pt(14)
        # 5. Oddiy matn (bold, italic va boshqa formatlashni qo‘llab-quvvatlash)
        else:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            parts = re.split(r'(\*\*.*?\*\*|__.*?__|\*.*?\*|_.*?_|~~.*?~~|`.*?`)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**") or part.startswith("__") and part.endswith("__"):
                    run = para.add_run(part[2:-2].strip())
                    run.bold = True
                    run.font.size = Pt(14)
                elif part.startswith("*") and part.endswith("*") or part.startswith("_") and part.endswith("_"):
                    run = para.add_run(part[1:-1].strip())
                    run.italic = True
                    run.font.size = Pt(14)
                elif part.startswith("~~") and part.endswith("~~"):
                    run = para.add_run(part[2:-2].strip())
                    run.font.strike = True
                    run.font.size = Pt(14)
                elif part.startswith("`") and part.endswith("`"):
                    run = para.add_run(part[1:-1].strip())
                    run.font.name = "Courier New"
                    run.font.color.rgb = RGBColor(128, 128, 128)
                    run.font.size = Pt(14)
                else:
                    run = para.add_run(part.strip())
                    run.bold = is_bold
                    run.italic = is_italic
                    run.font.size = Pt(14)

    para.paragraph_format.space_after = Pt(8)


def generate_bob_1(fan_nomi: str, mavzu: str) -> str:
    # Bo‘limlar uchun ma'lumotlarni generatsiya qilish
    sections = [
        {"title": "1.1. Yurak urishining mexanizmlari",
         "prompt": "Yurak urishining mexanizmlari haqida ilmiy ma'lumot ber, 2000 so‘zdan iborat bo‘lsin, O‘zbek tilida, ilmiy uslubda, plagiatdan xoli bo‘lsin. Hech qanday Kirish va Xulosa qismlari bo'lmasin."},
        {"title": "1.2. Elektrofiziologik jarayonlar",
         "prompt": "Yurak urishidagi elektrofiziologik jarayonlar haqida ilmiy ma'lumot ber, 2000 so‘zdan iborat bo‘lsin, O‘zbek tilida, ilmiy uslubda, plagiatdan xoli bo‘lsin. Hech qanday Kirish va Xulosa qismlari bo'lmasin."},
        {"title": "1.3. Yurak ritmining nazorati va boshqaruvi",
         "prompt": "Yurak ritmining nazorati va boshqaruvi haqida ilmiy ma'lumot ber, 2000 so‘zdan iborat bo‘lsin, O‘zbek tilida, ilmiy uslubda, plagiatdan xoli bo‘lsin. Hech qanday Kirish va Xulosa qismlari bo'lmasin."},
    ]

    # DOCX fayl yaratish
    document = Document()

    # I BOB sarlavhasi (markazda, qalin)
    add_formatted_paragraph(document, "I. BOB. YURAK URISHINING FIZIOLOGIYASI", is_heading=True, is_center=True)

    # Har bir bo‘limni generatsiya qilish va qo‘shish
    for section in sections:
        # Bo‘lim sarlavhasi (chapda, qalin)
        add_formatted_paragraph(document, section["title"], is_heading=True, is_center=False)

        # OpenAI'dan ma'lumot olish
        try:
            response = client.chat.completions.create(
                model="gpt-4.1-nano",  # To‘g‘ri modeldan foydalanamiz
                messages=[{"role": "user", "content": section["prompt"]}],
                temperature=0.7,
                max_tokens=6000,  # 2000 so‘z taxminan 6000 token atrofida bo‘ladi
            )

            section_text = response.choices[0].message.content.strip()
        except Exception as e:
            raise Exception(f"OpenAI API xatosi: {str(e)}")

        # Bo‘lim matnini qo‘shish
        lines = section_text.split('\n')
        for line in lines:
            if not line.strip():
                continue
            add_formatted_paragraph(document, line, is_heading=False)

    # Faylni saqlash
    try:
        save_path = f"generated_docs/bob_1_{sanitize_filename(mavzu)}.docx"
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        document.save(save_path)
    except Exception as e:
        raise Exception(f"Fayl saqlashda xato: {str(e)}")

    return save_path


if __name__ == '__main__':
    try:
        natija = generate_bob_1('Tibbiyot', "Yurak urishi va uni to'xtab qolishi")
        print(f"Fayl muvaffaqiyatli yaratildi: {natija}")
    except Exception as e:
        print(f"Xatolik yuz berdi: {str(e)}")
