from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os


def copy_paragraph_formatting(from_para, to_para):
    """
    Paragrafning formatlashini (masalan, tekislanish, bo'shliqlar) ko'chirish.
    """
    to_para.alignment = from_para.alignment
    to_para.paragraph_format.space_before = from_para.paragraph_format.space_before
    to_para.paragraph_format.space_after = from_para.paragraph_format.space_after
    to_para.paragraph_format.line_spacing = from_para.paragraph_format.line_spacing


def copy_run_formatting(from_run, to_run):
    """
    Run (matn qismi) formatlashini (masalan, qalin, shrift o'lchami) ko'chirish.
    """
    to_run.bold = from_run.bold
    to_run.italic = from_run.italic
    to_run.underline = from_run.underline
    to_run.font.size = from_run.font.size
    to_run.font.name = from_run.font.name
    to_run.font.color.rgb = from_run.font.color.rgb


def merge_docx_files(docx_files: list, output_path: str) -> str:
    """
    Bir nechta docx fayllarni birlashtirib, bitta docx fayl sifatida saqlaydi.
    Har bir docx faylning tarkibi yangi sahifadan boshlanadi va formatlash saqlanadi.

    Args:
        docx_files (list): Birlashtiriladigan docx fayllar yo'llari ro'yxati.
        output_path (str): Natija faylning saqlanadigan yo'li.

    Returns:
        str: Birlashtirilgan docx faylning yo'li.
    """
    # Yangi docx hujjat yaratish
    merged_document = Document()

    for i, docx_path in enumerate(docx_files):
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Fayl topilmadi: {docx_path}")

        # Agar bu birinchi docx fayl bo'lmasa, yangi sahifa qo'shish
        if i > 0:
            merged_document.add_page_break()

        # Docx faylni o'qish
        doc = Document(docx_path)

        # Har bir paragrafni ko'chirish
        for para in doc.paragraphs:
            if not para.text.strip():  # Bo'sh paragraflarni o'tkazib yuborish
                continue

            # Yangi paragraf yaratish
            new_para = merged_document.add_paragraph()

            # Paragraf formatlashini ko'chirish
            copy_paragraph_formatting(para, new_para)

            # Har bir run (matn qismi) ni ko'chirish
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                copy_run_formatting(run, new_run)

    # Birlashtirilgan faylni saqlash
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    merged_document.save(output_path)
    return output_path
