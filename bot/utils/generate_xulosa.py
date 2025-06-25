from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
import os
import re

# Atrof-muhit o‘zgaruvchilarini yuklash
load_dotenv()

# OpenAI mijozini boshlash
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# "Xulosa" sarlavhasini tilga moslashtirish uchun lug‘at
CONCLUSION_TITLES = {
    "o'zbek tili": "Xulosa",
    "rus tili": "Заключение",
    "ingliz tili": "Conclusion"
}

# Prompt uchun tilga mos shablonlar
PROMPT_TEMPLATES = {
    "o'zbek tili": {
        "conclusion": "Xulosa",
        "example_start": "Ushbu kurs ishi davomida {mavzu} bo‘yicha keng ko‘lamda o‘rganildi.",
        "example_follow": "Tahlillar natijasida aniqlanishicha, yurak urishining buzilishi ko‘pincha tashqi omillar va ichki patologik holatlar bilan bog‘liq bo‘lib, ularni oldindan aniqlash va bartaraf etish muhim ahamiyatga ega. Shu bilan birga, ushbu tadqiqot kelajakda davom ettirilishi va yanada chuqur tahlil qilinishi lozim."
    },
    "rus tili": {
        "conclusion": "Заключение",
        "example_start": "В ходе данной курсовой работы {mavzu} были всесторонне изучены.",
        "example_follow": "Анализ показал, что нарушения сердечного ритма часто связаны с внешними факторами и внутренними патологическими состояниями, что подчеркивает важность их своевременного выявления и устранения. Вместе с тем, данное исследование требует дальнейшего продолжения и более глубокого анализа."
    },
    "ingliz tili": {
        "conclusion": "Conclusion",
        "example_start": "Throughout this course work, {mavzu} were thoroughly investigated.",
        "example_follow": "The analysis revealed that disruptions in heart rhythm are often associated with external factors and internal pathological conditions, highlighting the importance of their early detection and mitigation. At the same time, this study needs to be continued and analyzed in greater depth in the future."
    }
}


def sanitize_filename(filename: str) -> str:
    """Fayl nomini maxsus belgilardan tozalash."""
    return re.sub(r'[^\w\s-]', '', filename).strip().replace(' ', '_')


def add_formatted_paragraph(document: Document, text: str, is_heading: bool = False) -> None:
    """Matnni formatlangan holda DOCX fayliga qo‘shish."""
    para = document.add_paragraph()

    # Matnni formatlash
    if is_heading:
        # Sarlavha uchun markazda va qalin
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run(text.strip())
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"
    else:
        # Oddiy matn uchun ikki tomondan tekislangan va qalin emas
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        run = para.add_run(text.strip())
        run.bold = False
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"

    para.paragraph_format.space_after = Pt(8)


def generate_xulosa(fan_nomi: str, mavzu: str, til: str, sahifa: int, chapter_1_sections: dict,
                    chapter_2_sections: dict) -> str:
    """
    Xulosa qismini generatsiya qilish va DOCX fayl sifatida saqlash.

    Args:
        fan_nomi (str): Fanning nomi.
        mavzu (str): Kurs ishi mavzusi.
        til (str): Til ("o'zbek tili", "rus tili", "ingliz tili").
        sahifa (int): Umumiy sahifalar soni (xulosa uzunligini hisoblash uchun).
        chapter_1_sections (dict): I bob bo'limlari.
        chapter_2_sections (dict): II bob bo'limlari.

    Returns:
        str: Saqlangan DOCX fayl yo'li.
    """
    # Tilga mos shablonni olish
    template = PROMPT_TEMPLATES.get(til.lower(), PROMPT_TEMPLATES["o'zbek tili"])  # Default: O‘zbek tili

    # Xulosa uzunligini hisoblash: umumiy sahifalar sonining taxminan 10-15% (1 sahifa ~ 300-400 so‘z)
    xulosa_sahifa = max(1, sahifa // 10)  # Xulosa kamida 1 sahifa bo‘ladi
    target_word_count = xulosa_sahifa * 350  # 1 sahifa ~ 350 so‘z
    target_tokens = target_word_count * 3  # 1 so‘z ~ 3 token

    # I va II bob bo‘limlarini promptga qo‘shish uchun tayyorlash
    chapter_1_info = "\n".join(
        [f"{key} {value}" for key, value in chapter_1_sections.items() if key != "chapter_title"])
    chapter_2_info = "\n".join(
        [f"{key} {value}" for key, value in chapter_2_sections.items() if key != "chapter_title"])

    # Dinamik example yaratish
    example = f"{template['example_start'].format(mavzu=mavzu)}\n{template['example_follow']}"

    # 1. OpenAI'dan xulosa qismi uchun matn olish
    prompt = f"""
    '{fan_nomi}' fanidan '{mavzu}' mavzusida kurs ishi uchun '{template['conclusion']}' qismini yozib ber. Quyidagi talablarga rioya qil, {til} tilida:
    - Xulosa matni sarlavhasiz bo‘lsin.
    - Xulosa matni taxminan {target_word_count} so‘zdan iborat bo‘lsin.
    - Matn kurs ishining umumiy natijalari, muhim topilmalar va yakuniy fikrlarni o‘z ichiga olishi kerak.
    - Matn ilmiy uslubda, aniq, ravon va plagiatdan xoli bo‘lsin.
    - Quyidagi bo‘limlar asosida xulosa yoz:
      I bob: {chapter_1_sections.get('chapter_title', 'Nomalum')}
      {chapter_1_info}
      II bob: {chapter_2_sections.get('chapter_title', 'Nomalum')}
      {chapter_2_info}
    - {til} tilida yozing.
    Times New Roman, 14 pt.
    Namuna sifatida quyidagi uslubdan foydalaning:
    {example}
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4.1-nano",  # Agar model ishlamasa, "gpt-3.5-turbo"ga o‘zgartirish mumkin
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=target_tokens
        )
        xulosa_text = response.choices[0].message.content.strip()
    except Exception as e:
        raise Exception(f"OpenAI API xatosi: {str(e)}")

    # 2. DOCX fayl yaratish
    document = Document()

    # Xulosa sarlavhasi (markazda, qalin, tilga mos)
    conclusion_title = CONCLUSION_TITLES.get(til.lower(), "Xulosa")  # Agar til topilmasa, default O‘zbek tili
    add_formatted_paragraph(document, conclusion_title, is_heading=True)

    # Xulosa matnini qo‘shish
    lines = xulosa_text.split('\n')
    for line in lines:
        if not line.strip():
            continue
        add_formatted_paragraph(document, line, is_heading=False)

    # 3. Saqlash
    try:
        sanitized_mavzu = sanitize_filename(mavzu)
        save_path = f"generated_docs/xulosa_{sanitized_mavzu}_{til.lower().replace(' ', '_')}.docx"
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        document.save(save_path)
    except Exception as e:
        raise Exception(f"Fayl saqlashda xato: {str(e)}")

    return save_path
